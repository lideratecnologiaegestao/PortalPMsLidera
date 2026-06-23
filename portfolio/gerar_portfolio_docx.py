# -*- coding: utf-8 -*-
"""
Gera o portfólio comercial do "Portal de Prefeitura" em .docx, com gráficos.
Dependências: python-docx, matplotlib, Pillow.
Uso: python gerar_portfolio_docx.py
Saída: portfolio/Portfolio-Portal-Prefeitura.docx + portfolio/assets/*.png
"""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------------- paleta
AZUL      = "#1351B4"
AZUL_ESC  = "#0C326F"
AZUL_CLARO= "#5992ED"
VERDE     = "#168821"
AMARELO   = "#FFCD07"
VERMELHO  = "#E52207"
CINZA_BG  = "#F2F5FA"
CINZA_LN  = "#D9E0EE"
TEXTO     = "#1F2933"

HEX = lambda c: RGBColor.from_string(c.lstrip("#"))

BASE   = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(BASE, "assets")
os.makedirs(ASSETS, exist_ok=True)

plt.rcParams.update({
    "font.family": "DejaVu Sans",
    "font.size": 12,
    "svg.fonttype": "none",
})

# ======================================================================
#  GRÁFICOS (matplotlib -> PNG)
# ======================================================================
def _box(ax, x, y, w, h, text, fc, tc="white", fs=11, bold=True, ec=None, lw=0):
    ax.add_patch(FancyBboxPatch((x, y), w, h,
                 boxstyle="round,pad=0.012,rounding_size=0.03",
                 fc=fc, ec=ec or fc, lw=lw, mutation_aspect=1))
    ax.text(x + w/2, y + h/2, text, ha="center", va="center",
            color=tc, fontsize=fs, fontweight="bold" if bold else "normal",
            wrap=True, zorder=5)

def _arrow(ax, p1, p2, color=AZUL_ESC):
    ax.add_patch(FancyArrowPatch(p1, p2, arrowstyle="-|>", mutation_scale=14,
                 lw=1.6, color=color, shrinkA=2, shrinkB=2))

def fig_banner():
    fig, ax = plt.subplots(figsize=(11, 4.0), dpi=200)
    grad = np.linspace(0, 1, 256).reshape(1, -1)
    ax.imshow(grad, extent=[0, 10, 0, 4], aspect="auto",
              cmap=matplotlib.colors.LinearSegmentedColormap.from_list(
                  "g", [AZUL_ESC, AZUL]), zorder=0)
    # formas decorativas
    for cx, r, a in [(8.7, 2.2, .10), (9.4, 1.3, .14), (1.0, 2.6, .07)]:
        ax.add_patch(plt.Circle((cx, 2), r, color="white", alpha=a, zorder=1))
    ax.add_patch(plt.Rectangle((0, 0), 10, .18, color=AMARELO, zorder=3))
    ax.add_patch(plt.Rectangle((0, 0), 3.3, .18, color=VERDE, zorder=4))
    ax.text(0.5, 2.62, "PORTAL DE PREFEITURA", color="white", fontsize=33,
            fontweight="bold", va="center", zorder=5)
    ax.text(0.55, 1.75, "A plataforma digital completa, acessível e em",
            color="white", fontsize=15, va="center", zorder=5, alpha=.95)
    ax.text(0.55, 1.32, "conformidade com a lei para o seu município.",
            color="white", fontsize=15, va="center", zorder=5, alpha=.95)
    ax.text(0.55, 0.55, "LIDERA  ·  Tecnologia e Gestão", color=AMARELO,
            fontsize=13, fontweight="bold", va="center", zorder=5)
    ax.set_xlim(0, 10); ax.set_ylim(0, 4); ax.axis("off")
    p = os.path.join(ASSETS, "banner.png")
    fig.savefig(p, bbox_inches="tight", pad_inches=0); plt.close(fig); return p

def fig_arquitetura():
    fig, ax = plt.subplots(figsize=(10, 6.2), dpi=200)
    ax.set_xlim(0, 10); ax.set_ylim(0, 10); ax.axis("off")
    _box(ax, 3.3, 9.0, 3.4, 0.8, "Cidadão  ·  App do Cidadão", AZUL_ESC, fs=12)
    _box(ax, 3.3, 7.9, 3.4, 0.7, "Cloudflare  ·  WAF  ·  TLS", "#6B7A99", fs=11)
    _box(ax, 2.7, 6.9, 4.6, 0.6, "Nginx · roteamento por domínio", "#8895AE", fs=10)
    _box(ax, 0.8, 5.5, 4.0, 0.9, "Portal Web\nNext.js (SSR/ISR)", AZUL, fs=11)
    _box(ax, 5.2, 5.5, 4.0, 0.9, "API\nNestJS (monólito modular)", AZUL, fs=11)
    # camada de dados
    _box(ax, 0.5, 3.6, 2.7, 0.95, "PostgreSQL 16\n+ PostGIS  ·  RLS", VERDE, fs=10)
    _box(ax, 3.45, 3.6, 2.6, 0.95, "Redis\n+ BullMQ (filas)", VERDE, fs=10)
    _box(ax, 6.3, 3.6, 2.9, 0.95, "MinIO / S3\n(arquivos)", VERDE, fs=10)
    _box(ax, 1.7, 2.2, 3.0, 0.8, "Workers assíncronos\n(SLA · OCR · notificações)", "#0E6B1A", fs=9)
    _box(ax, 5.0, 2.2, 3.2, 0.8, "IA (Claude)\ntriagem · RAG · OCR", "#0E6B1A", fs=9)
    for a, b in [((5,9.0),(5,8.6)), ((5,7.9),(5,7.5)), ((5,6.9),(5,6.4)),
                 ((4.3,6.9),(2.8,6.4)), ((5.7,6.9),(7.2,6.4)),
                 ((2.8,5.5),(1.8,4.55)), ((7.2,5.5),(7.7,4.55)),
                 ((7.2,5.5),(4.7,4.55)), ((7.2,5.5),(3.2,4.0))]:
        _arrow(ax, a, b)
    # banner camadas de segurança
    ax.add_patch(FancyBboxPatch((0.3, 0.5), 9.4, 1.05,
                 boxstyle="round,pad=0.02,rounding_size=0.05",
                 fc=AMARELO, ec=AMARELO))
    ax.text(5, 1.27, "DUAS CAMADAS DE SEGURANÇA INDEPENDENTES", ha="center",
            color=AZUL_ESC, fontsize=11, fontweight="bold")
    ax.text(2.6, 0.83, "RBAC — o que cada um PODE FAZER", ha="center",
            color=TEXTO, fontsize=9.5)
    ax.text(7.4, 0.83, "RLS — o que cada um PODE VER", ha="center",
            color=TEXTO, fontsize=9.5)
    p = os.path.join(ASSETS, "arquitetura.png")
    fig.savefig(p, bbox_inches="tight", pad_inches=0.1); plt.close(fig); return p

def fig_multitenant():
    fig, ax = plt.subplots(figsize=(10, 4.8), dpi=200)
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis("off")
    _box(ax, 3.2, 4.7, 3.6, 1.0, "UMA plataforma\nUMA infraestrutura", AZUL_ESC, fs=13)
    nomes = ["Prefeitura A", "Prefeitura B", "Prefeitura C", "Prefeitura D", "Prefeitura N"]
    cores = [AZUL, VERDE, "#B3531D", "#6f42c1", "#155F8A"]
    xs = np.linspace(0.4, 7.7, 5)
    for x, n, c in zip(xs, nomes, cores):
        _box(ax, x, 1.6, 1.75, 1.15, n + "\n", c, fs=10)
        ax.text(x + 0.875, 1.95, "domínio · tema · conteúdo", ha="center",
                va="center", color="white", fontsize=6.6)
        _arrow(ax, (5, 4.7), (x + 0.875, 2.75))
    ax.text(5, 0.7, "Atualizou uma vez  →  todas as prefeituras recebem.  "
            "Dados 100% isolados por RLS.", ha="center", color=TEXTO,
            fontsize=10.5, fontweight="bold")
    p = os.path.join(ASSETS, "multitenant.png")
    fig.savefig(p, bbox_inches="tight", pad_inches=0.1); plt.close(fig); return p

def fig_conformidade():
    itens = ["gov.br — Login Único", "PNTP / Atricon", "WCAG 2.1 AA / eMAG",
             "LGPD (Lei 13.709/18)", "LC 131 + LRF (Transparência)",
             "Lei 13.460/17 (Ouvidoria)", "LAI (Lei 12.527/11)"]
    fig, ax = plt.subplots(figsize=(10, 4.6), dpi=200)
    y = np.arange(len(itens))
    ax.barh(y, [100]*len(itens), color=AZUL, height=0.62, zorder=3)
    ax.barh(y, [100]*len(itens), color=CINZA_BG, height=0.62, zorder=1)
    ax.barh(y, [100]*len(itens), color=AZUL, height=0.62, zorder=3)
    for yi in y:
        ax.text(98, yi, "✓ Atende", ha="right", va="center", color="white",
                fontsize=11, fontweight="bold", zorder=4)
    ax.set_yticks(y); ax.set_yticklabels(itens, fontsize=11, color=TEXTO)
    ax.set_xticks([]); ax.set_xlim(0, 100)
    for s in ax.spines.values(): s.set_visible(False)
    ax.set_title("Conformidade legal nativa", fontsize=14, color=AZUL_ESC,
                 fontweight="bold", loc="left", pad=12)
    p = os.path.join(ASSETS, "conformidade.png")
    fig.savefig(p, bbox_inches="tight", pad_inches=0.1); plt.close(fig); return p

def fig_pntp():
    fig, (axd, axb) = plt.subplots(1, 2, figsize=(10, 4.2), dpi=200,
                                   gridspec_kw={"width_ratios": [1, 1.4]})
    # donut — 100% dos critérios PNTP atendidos (selo Diamante)
    axd.pie([100], colors=[VERDE], startangle=90,
            counterclock=False, wedgeprops=dict(width=0.32))
    axd.text(0, 0.20, "100%", ha="center", fontsize=22, fontweight="bold",
             color=VERDE)
    axd.text(0, -0.16, "critérios PNTP", ha="center", fontsize=9.5, color=TEXTO)
    axd.text(0, -0.40, "SELO DIAMANTE", ha="center", fontsize=9.5,
             fontweight="bold", color=AZUL_ESC)
    axd.set_aspect("equal")
    # barras de dimensões com mais peso
    dims = ["Receita", "Despesa", "Recursos\nHumanos", "Licitações", "Contratos", "Obras"]
    pesos = [4, 4, 3, 3, 3, 2]
    cores = [AZUL if p == 4 else (AZUL_CLARO if p == 3 else "#9DBDF0") for p in pesos]
    axb.bar(range(len(dims)), pesos, color=cores, zorder=3, width=0.66)
    axb.set_xticks(range(len(dims))); axb.set_xticklabels(dims, fontsize=9, color=TEXTO)
    axb.set_yticks([1,2,3,4]); axb.set_ylim(0, 4.6)
    axb.set_ylabel("Peso da dimensão", fontsize=10, color=TEXTO)
    for i, p in enumerate(pesos):
        axb.text(i, p+0.12, str(p), ha="center", fontsize=10, fontweight="bold", color=AZUL_ESC)
    for s in ["top", "right"]: axb.spines[s].set_visible(False)
    axb.set_title("Dimensões de maior peso (Transparência)", fontsize=11,
                  color=AZUL_ESC, fontweight="bold", loc="left")
    p = os.path.join(ASSETS, "pntp.png")
    fig.savefig(p, bbox_inches="tight", pad_inches=0.12); plt.close(fig); return p

def fig_roadmap():
    fases = [("Fundação\n(RLS, núcleo, e-SIC)", True),
             ("Transparência\n(ETL, dados abertos)", True),
             ("Serviços + CMS", True),
             ("Diário Oficial\n(ICP-Brasil)", True),
             ("Inteligência\nArtificial", True),
             ("App do Cidadão", False),
             ("Escala /\nMulti-região", False)]
    fig, ax = plt.subplots(figsize=(10, 2.7), dpi=200)
    ax.set_xlim(0, len(fases)); ax.set_ylim(0, 2); ax.axis("off")
    ax.plot([0.3, len(fases)-0.3], [1.35, 1.35], color=CINZA_LN, lw=3, zorder=1)
    for i, (txt, done) in enumerate(fases):
        x = i + 0.5
        col = VERDE if done else AZUL
        ax.scatter([x], [1.35], s=420, color=col, zorder=3, edgecolors="white", lw=2)
        ax.text(x, 1.35, "✓" if done else str(i), ha="center", va="center",
                color="white", fontsize=11, fontweight="bold", zorder=4)
        ax.text(x, 0.75, txt, ha="center", va="top", fontsize=8.6, color=TEXTO)
        if done:
            ax.text(x, 1.78, "CONCLUÍDA", ha="center", fontsize=8,
                    color=VERDE, fontweight="bold")
    p = os.path.join(ASSETS, "roadmap.png")
    fig.savefig(p, bbox_inches="tight", pad_inches=0.1); plt.close(fig); return p

def fig_kpis():
    kpis = [("100% PNTP", "Selo Diamante\nem transparência"),
            ("1 → N", "uma plataforma,\ninfinitas prefeituras"),
            ("20+", "módulos\nintegrados"),
            ("IA + RAG", "busca semântica\ne respostas da base oficial"),
            ("WCAG 2.1 AA", "acessibilidade\npor padrão")]
    fig, axes = plt.subplots(1, 5, figsize=(11, 2.3), dpi=200)
    for ax, (n, d) in zip(axes, kpis):
        ax.axis("off")
        ax.add_patch(FancyBboxPatch((0.04, 0.08), 0.92, 0.84,
                     boxstyle="round,pad=0.01,rounding_size=0.06",
                     transform=ax.transAxes, fc=CINZA_BG, ec=CINZA_LN, lw=1))
        ax.text(0.5, 0.66, n, ha="center", va="center", transform=ax.transAxes,
                fontsize=16, fontweight="bold", color=AZUL)
        ax.text(0.5, 0.30, d, ha="center", va="center", transform=ax.transAxes,
                fontsize=8.6, color=TEXTO)
    p = os.path.join(ASSETS, "kpis.png")
    fig.savefig(p, bbox_inches="tight", pad_inches=0.05); plt.close(fig); return p

print("Gerando gráficos...")
IMG = {
    "banner": fig_banner(), "arquitetura": fig_arquitetura(),
    "multitenant": fig_multitenant(), "conformidade": fig_conformidade(),
    "pntp": fig_pntp(), "roadmap": fig_roadmap(), "kpis": fig_kpis(),
}

# ======================================================================
#  DOCX
# ======================================================================
def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), color.lstrip("#")); tcPr.append(sh)

def set_cell_text(cell, text, bold=False, color=None, size=10, white=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align: p.alignment = align
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if white: r.font.color.rgb = HEX("#FFFFFF")
    elif color: r.font.color.rgb = HEX(color)

def no_table_borders(table):  # subtle borders
    tbl = table._tbl; tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"),"single")
        e.set(qn("w:sz"),"4"); e.set(qn("w:color"),"D9E0EE")
        borders.append(e)
    tblPr.append(borders)

doc = Document()
# A4 + margens
sec = doc.sections[0]
sec.page_height = Emu(int(11.69*914400)); sec.page_width = Emu(int(8.27*914400))
for m in ("top","bottom"): setattr(sec, f"{m}_margin", Inches(0.7))
for m in ("left","right"): setattr(sec, f"{m}_margin", Inches(0.85))
CW = Inches(6.5)  # largura de conteúdo

# estilos base
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)
st.font.color.rgb = HEX(TEXTO)
for h, sz, col in [("Heading 1",17,AZUL_ESC),("Heading 2",13.5,AZUL),("Heading 3",11.5,AZUL_ESC)]:
    s = doc.styles[h]; s.font.size = Pt(sz); s.font.color.rgb = HEX(col)
    s.font.name = "Calibri"; s.font.bold = True

def h1(t):
    doc.add_paragraph()
    p = doc.add_heading(t, level=1)
    # faixa fina abaixo
    return p
def h2(t): return doc.add_heading(t, level=2)
def para(t, size=10.5, italic=False, color=None, bold=False, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(t); r.font.size = Pt(size); r.italic = italic; r.bold = bold
    if color: r.font.color.rgb = HEX(color)
    return p
def bullets(items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it); r.font.size = Pt(size)
def img(key, width=CW, caption=None):
    doc.add_picture(IMG[key], width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption); r.italic = True; r.font.size = Pt(8.5)
        r.font.color.rgb = HEX("#6B7A99")

# ---------------- CAPA ----------------
doc.add_picture(IMG["banner"], width=CW)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
r = p.add_run("Portfólio Institucional & Comercial"); r.bold = True
r.font.size = Pt(15); r.font.color.rgb = HEX(AZUL_ESC)
para("Plataforma SaaS multi-tenant para a gestão pública municipal",
     size=11.5, italic=True, color="#444455", after=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
img("kpis", width=CW)
# logo oficial da Lidera (PNG rasterizado do SVG vetorial)
_logo = os.path.join(ASSETS, "logo_lidera.png")
if os.path.exists(_logo):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.add_run().add_picture(_logo, width=Inches(2.5))
else:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("Lidera Tecnologia e Gestão"); r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = HEX(AZUL)
for t in ["www.lideratecnologia.com.br   ·   comercial@lideratecnologia.com.br",
          "Demonstração ao vivo:  https://prefeitura.lidera.app.br"]:
    q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = q.add_run(t); rr.font.size = Pt(9.5); rr.font.color.rgb = HEX("#556677")
doc.add_page_break()

# ---------------- 1. O DESAFIO / A SOLUÇÃO ----------------
h1("O desafio das prefeituras — e a nossa solução")
para("Pequenos e médios municípios enfrentam, ao mesmo tempo, exigências legais crescentes "
     "e orçamentos de TI limitados. O resultado é conhecido: sites desatualizados, portais "
     "de transparência incompletos, prazos legais de ouvidoria e e-SIC perdidos, e barreiras "
     "de acessibilidade que excluem o cidadão.")
h2("As dores mais comuns")
bullets([
 "Transparência e PNTP: dificuldade de atingir e manter o selo da Atricon, com dados dispersos e desatualizados.",
 "Prazos legais: e-SIC (LAI) e Ouvidoria (Lei 13.460) exigem controle de prazos que planilhas não dão conta.",
 "Acessibilidade: a lei exige WCAG/eMAG e VLibras — a maioria dos sites não cumpre.",
 "Custo e dependência: cada sistema (site, transparência, diário, ouvidoria) é um contrato e um fornecedor diferente.",
 "LGPD: tratamento de dados do cidadão sem base legal, sem direitos do titular e sem registro de incidentes.",
])
h2("A solução: uma única plataforma, completa e em conformidade")
para("O Portal de Prefeitura reúne site institucional, transparência, ouvidoria/e-SIC, diário "
     "oficial, serviços ao cidadão, app mobile e inteligência artificial em um só produto — "
     "com a legislação embutida e a acessibilidade garantida por padrão.")

# ---------------- 2. VISÃO GERAL / ARQUITETURA ----------------
h1("Visão geral da plataforma")
para("Construída sobre tecnologias modernas e abertas (NestJS, Next.js, PostgreSQL/PostGIS, "
     "Redis, n8n e IA), a plataforma é segura por projeto: o frontend e o app falam somente "
     "com a API (gateway único), e os dados de cada prefeitura ficam isolados no banco por "
     "Row Level Security (RLS).")
img("arquitetura", width=Inches(5.7),
    caption="Arquitetura em camadas, com as duas camadas de segurança (RBAC + RLS).")

# ---------------- 3. MULTI-TENANT ----------------
h1("Modelo SaaS multi-tenant")
para("Uma única base de código e uma única infraestrutura servem N prefeituras. Cada município "
     "tem domínio, identidade visual (white-label) e conteúdo próprios — mas todos compartilham "
     "o mesmo motor, sempre atualizado.")
img("multitenant", width=Inches(6.2),
    caption="Uma plataforma → muitas prefeituras. Custo compartilhado, evolução simultânea.")
h2("Por que isso importa para o seu município")
bullets([
 "Custo compartilhado: a infraestrutura é diluída entre todas as prefeituras atendidas.",
 "Atualização simultânea: novas funções e correções legais chegam a todos sem projeto novo.",
 "Isolamento comprovado: os dados de uma prefeitura nunca são visíveis para outra (RLS).",
 "Identidade própria: cores, logo, domínio e conteúdo são exclusivos de cada município.",
])
doc.add_page_break()

# ---------------- 4. MÓDULOS ----------------
h1("Catálogo de módulos")
para("Mais de 20 módulos integrados, organizados em seis grandes áreas. Tudo no mesmo login, "
     "no mesmo painel e com os mesmos dados.")
MOD = [
 ("Cidadão & Participação", AZUL, [
    "Ouvidoria + e-SIC (LAI e Lei 13.460) com máquina de estados, prazos/SLA, chat de tramitação e painel do ouvidor — sigilo da fonte garantido por duas camadas de controle",
    "Atendimento omnichannel: bot 24h que EXECUTA ações (abre manifestação com protocolo, consulta andamento, chama atendente) + WhatsApp multi-provedor com redundância + caixa unificada",
    "Enquetes (consulta pública anônima) e Construtor de formulários (drag-drop, captcha anti-robô Cloudflare Turnstile, exportação CSV/XML/Excel)",
    "App do Cidadão: chamados georreferenciados (buracos, terrenos, animais) com foto e GPS",
    "Comentários em notícias com moderação automática por IA (barra ofensas e código malicioso, preserva crítica legítima)",
 ]),
 ("Transparência & Prestação de Contas", VERDE, [
    "Portal da Transparência (LC 131/LRF): receitas, despesas, licitações, contratos e folha, com dados abertos (CSV/JSON)",
    "Importação automática do APLIC/TCE-MT",
    "Conformidade PNTP/Atricon — 100% dos critérios atendidos (Selo Diamante)",
    "Diário Oficial 2.0: matérias estruturadas, busca full-text em português e assinatura ICP-Brasil",
 ]),
 ("Conteúdo & Comunicação (CMS)", "#B3531D", [
    "Construtor de páginas drag-drop (11 tipos de blocos, templates, versões/backup, SEO por página)",
    "Notícias, Secretarias (8 seções), Galeria (vídeos mp4/YouTube)",
    "Cadastro de Documentos (Leis, Decretos… com tipos TCE-MT e contador de downloads)",
    "Estrutura organizacional automática e Carta de Serviços",
 ]),
 ("Gestão & Administração", "#6f42c1", [
    "Gerenciador multi-tenant (cadastro de prefeituras) e Configurações por entidade",
    "Console da Plataforma (configuração global)",
    "Usuários, Grupos e Sessões: permissões granulares e sessões online com revogação",
    "Chat interno entre servidores e Home configurável (Acesso Rápido)",
 ]),
 ("Inteligência Artificial", "#155F8A", [
    "Chatbot com busca híbrida (FTS + semântica/embeddings + reranking Voyage): responde a partir da base oficial combinando palavra e significado, lê PDFs digitalizados via OCR — isolado por município",
    "Base de conhecimento treinável pelo gestor: além de FAQs, o servidor cadastra Artigos livres (normas, eventos, materiais educativos) que a IA passa a usar — sem programação",
    "Bot agêntico na ouvidoria: executa por linguagem natural (abre manifestação e devolve protocolo, consulta andamento, chama atendente humano)",
    "Assistente interno que ensina servidores a usar o sistema + Manual do Sistema integrado ao painel",
    "Triagem automática de manifestações (sempre com revisão humana) e IA fiscal sobre a transparência",
 ]),
 ("Conformidade & Segurança", VERMELHO, [
    "LGPD self-service: direitos do titular, portabilidade, anonimização, incidentes e geração de documentos (prefeitura = controladora / Lidera = operadora)",
    "Sigilo multicamadas (RBAC + RLS por papel): denúncias da ouvidoria visíveis só ao ouvidor e equipe — nem o administrador da prefeitura acessa",
    "Autocadastro de usuário + solicitação de elevação por cargo/lotação; papéis sensíveis (ouvidor) aprovados pela operadora; termo de sigilo no primeiro acesso",
    "Escopo por secretaria: servidor lotado numa área só edita o conteúdo da sua secretaria",
    "Proteção anti-robô Cloudflare Turnstile em logins e formulários + RBAC + RLS + WCAG 2.1 AA / eMAG + VLibras",
 ]),
]
for titulo, cor, itens in MOD:
    tb = doc.add_table(rows=1, cols=1); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(tb)
    hd = tb.rows[0].cells[0]; shade(hd, cor)
    set_cell_text(hd, "  " + titulo, bold=True, white=True, size=11)
    bc = tb.add_row().cells[0]
    bc.text = ""
    for it in itens:
        pp = bc.add_paragraph(style="List Bullet")
        rr = pp.add_run(it); rr.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
doc.add_page_break()

# ---------------- 5. CONFORMIDADE ----------------
h1("Conformidade legal nativa")
para("A legislação não é um módulo à parte: ela está no núcleo do produto. Os prazos legais já "
     "vêm configurados e o sistema alerta e vence automaticamente.")
img("conformidade", width=Inches(6.2))
tb = doc.add_table(rows=1, cols=2); no_table_borders(tb); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_text(tb.rows[0].cells[0], "Norma", bold=True, white=True); shade(tb.rows[0].cells[0], AZUL)
set_cell_text(tb.rows[0].cells[1], "Como o portal atende", bold=True, white=True); shade(tb.rows[0].cells[1], AZUL)
for norma, como in [
 ("LAI — Lei 12.527/2011", "e-SIC com prazo de 20+10 dias e instâncias de recurso na máquina de estados"),
 ("Lei 13.460/2017", "Ouvidoria com prazo de 30+30 dias e tipos de manifestação"),
 ("LC 131/2009 + LRF", "Transparência ativa de receitas, despesas, licitações e contratos + dados abertos"),
 ("LGPD — Lei 13.709/2018", "Bases legais, direitos do titular self-service, anonimização e registro de incidentes"),
 ("WCAG 2.1 AA / eMAG", "Tema reprovado no contraste não é salvo; VLibras e Design System gov.br"),
 ("PNTP / Atricon", "100% dos critérios aplicáveis atendidos — Selo Diamante"),
 ("gov.br — Login Único", "Autenticação do cidadão por OIDC com selos de confiabilidade"),
]:
    row = tb.add_row().cells
    set_cell_text(row[0], norma, bold=True, size=9.5, color=AZUL_ESC)
    set_cell_text(row[1], como, size=9.5)

# ---------------- 6. TRANSPARÊNCIA / PNTP ----------------
h1("Transparência e PNTP — 100% dos critérios (Selo Diamante)")
para("O Programa Nacional de Transparência Pública (Atricon) avalia o portal por uma matriz de "
     "critérios ponderados. A plataforma atende a 100% dos critérios aplicáveis — nível máximo, "
     "o Selo Diamante — pontuando nos itens de maior peso (Receita, Despesa, Recursos Humanos, "
     "Licitações e Contratos) em disponibilidade, atualidade, série histórica, download em formato "
     "aberto e filtro de pesquisa.", bold=False)
para("Sua prefeitura entra no ar já alinhada à nota máxima da transparência pública.",
     bold=True, color=VERDE, after=8)
img("pntp", width=Inches(6.2))
doc.add_page_break()

# ---------------- 7. SEGURANÇA ----------------
h1("Segurança & Privacidade")
bullets([
 "Isolamento por RLS: cada consulta ao banco é restrita ao município; uma prefeitura nunca enxerga dados de outra.",
 "RBAC: papéis de domínio (administrador, gestor, ouvidor, servidor, cidadão) controlam cada ação.",
 "Sigilo jurídico da ouvidoria: denúncias acessíveis SOMENTE ao ouvidor e sua equipe — RBAC + RLS por papel; o administrador da prefeitura não tem acesso.",
 "Autocadastro + elevação de papel por autoridade: servidores declaram cargo e lotação; papéis sensíveis (ouvidor) exigem aprovação da operadora e termo de sigilo.",
 "Escopo por secretaria: cada servidor gerencia apenas o conteúdo da sua área.",
 "Proteção anti-robô Cloudflare Turnstile em logins e formulários — sem CAPTCHA tradicional.",
 "LGPD por projeto: minimização de dados, base legal por finalidade e logs de acesso a dados pessoais.",
 "Auditoria: toda ação sensível e toda falha de processamento são registradas.",
 "Borda protegida: nada exposto direto à internet — proxy reverso + WAF + TLS obrigatório.",
 "Login gov.br: identidade forte do cidadão, sem o município guardar senha.",
])

# ---------------- 8. APP ----------------
h1("App do Cidadão (mobile)")
para("Aplicativo em React Native (Expo) para o cidadão abrir chamados georreferenciados — "
     "buracos na via, terrenos abandonados, animais nas ruas — com foto e GPS. A plataforma "
     "detecta duplicados por proximidade e o cidadão acompanha o andamento e recebe notificações. "
     "Login via gov.br. Tema e identidade da prefeitura aplicados ao app (white-label).")

# ---------------- 8b. ATENDIMENTO OMNICHANNEL ----------------
h1("Atendimento Omnichannel com IA — 5 canais, uma caixa unificada")
para("O cidadão escolhe o canal. A prefeitura atende em todos — sem trocar de tela, sem perder historico. "
     "O bot de IA responde 24 horas, abre protocolos, consulta andamentos e escala para humano quando necessario.")
h2("5 canais integrados")
bullets([
 "Site / Widget: chat flutuante no proprio portal da prefeitura. Ativo por padrao, sem configuracao extra.",
 "WhatsApp API Oficial (Cloud API Meta): numero profissional homologado. Conformidade com editais que exigem API oficial.",
 "Instagram Direct: mensagens do perfil profissional da prefeitura chegam na mesma fila de atendimento.",
 "Facebook Messenger: pagina oficial da prefeitura integrada ao painel sem troca de tela.",
 "Telegram: bot via @BotFather. Webhook registrado automaticamente pelo sistema — sem configuracao manual.",
])
h2("Bot de IA 24 horas")
bullets([
 "Responde perguntas sobre servicos, documentos, horarios e requisitos com base no conteudo real da prefeitura.",
 "Abre manifestacoes de ouvidoria por linguagem natural e devolve protocolo + chave na hora.",
 "Consulta andamento de protocolos existentes sem o cidadao precisar ligar ou ir presencialmente.",
 "Reconhece quando precisa de atendente humano e transfere a conversa com notificacao imediata.",
 "Aprende com FAQs e Artigos cadastrados pelo gestor — sem codigo e sem TI.",
])
h2("Console do atendente — caixa unificada")
bullets([
 "Todos os canais (Widget, WhatsApp, Instagram, Messenger, Telegram) em uma unica tela.",
 "Fila de conversas com filtro por canal e por secretaria.",
 "Historico completo da conversa com o bot antes de cada transferencia.",
 "Distribuicao por secretaria: saude para equipe de saude, obras para obras.",
 "Relatorios de volume, tempo de resposta e canais mais usados.",
])
h2("Outros destaques")
bullets([
 "Multi-numero: 1 numero para toda a prefeitura ou varios por secretaria — gerenciados na mesma tela.",
 "Alerta de creditos de template WhatsApp (HSM): monitoramento do consumo com aviso de limite baixo.",
 "App instalavel (PWA): o cidadao instala o portal na tela do celular sem App Store; o chat funciona na PWA.",
 "API Oficial Meta — diferencial em editais que vedam versoes nao autorizadas.",
 "Chaves cifradas em repouso (AES-256-GCM). Isolamento por prefeitura via RLS. LGPD nativo.",
 "Incluido no Portal de Prefeitura sem custo adicional de modulo.",
])
doc.add_page_break()

# ---------------- 9. IA ----------------
h1("Inteligência Artificial aplicada")
para("A IA não é um chatbot genérico colado ao portal: está integrada ao núcleo, treinada com os dados reais da prefeitura e isolada por município. O gestor treina sem programação; o cidadão recebe respostas da base oficial.")
bullets([
 "Chatbot com busca híbrida avançada (FTS + semântica/embeddings com reranking Voyage): responde combinando busca por palavra e por significado — incluindo conteúdo de PDFs digitalizados via OCR (documentos antigos escaneados incluídos).",
 "Base de conhecimento treinável pelo gestor: além de FAQs, cadastre Artigos livres (normas de saúde, regimentos, materiais para alunos, eventos) e a IA começa a responder sobre o tema — sem programação.",
 "Bot agêntico no atendimento: além de responder, o bot EXECUTA por linguagem natural — abre manifestação de ouvidoria com protocolo+chave, consulta andamento e transfere para atendente humano.",
 "Assistente interno (\"Como faço X?\") que ensina os próprios servidores a usar o sistema, com Manual do Sistema integrado ao painel administrativo.",
 "Moderação inteligente de comentários em notícias: a IA barra ofensas, baixo calão e código malicioso sem censurar crítica ou opinião legítima — o que passa vai para aprovação humana.",
 "Triagem automática de manifestações de ouvidoria/e-SIC, sempre com revisão humana.",
 "IA fiscal sobre a transparência, com consulta estruturada aos dados contábeis do APLIC/TCE-MT.",
])

# ---------------- 10. PROVA ----------------
h1("Prova de entrega — caso real")
para("A Prefeitura de Barão de Melgaço (MT) teve seu site legado (Joomla) migrado integralmente "
     "para a plataforma, preservando o histórico e o SEO:", bold=False)
tb = doc.add_table(rows=1, cols=3); no_table_borders(tb); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
dados = [("455", "notícias migradas"), ("1.180", "documentos migrados"), ("13", "secretarias publicadas")]
for i,(n,d) in enumerate(dados):
    c = tb.rows[0].cells[i]; shade(c, CINZA_BG)
    c.text = ""
    p1 = c.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p1.add_run(n); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = HEX(AZUL)
    p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(d); r2.font.size = Pt(9.5); r2.font.color.rgb = HEX(TEXTO)
para("Além do conteúdo, foram migrados o institucional, a galeria e a Carta de Serviços, com "
     "redirecionamentos que preservam os links antigos indexados pelo Google.", after=4)
doc.add_page_break()

# ---------------- 11. IMPLANTAÇÃO ----------------
h1("Implantação flexível")
para("A plataforma roda onde a prefeitura precisar, com manuais e automação (Terraform) prontos:")
tb = doc.add_table(rows=1, cols=2); no_table_borders(tb); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_text(tb.rows[0].cells[0], "Modelo", bold=True, white=True); shade(tb.rows[0].cells[0], AZUL)
set_cell_text(tb.rows[0].cells[1], "O que oferece", bold=True, white=True); shade(tb.rows[0].cells[1], AZUL)
for a,b in [
 ("Nuvem gerenciada (Google Cloud / AWS)", "Provisionamento por Terraform, escala automática, backups e WAF"),
 ("On-premise (Windows / Linux / Docker)", "Roda no servidor da prefeitura ou do parceiro, com os mesmos recursos"),
 ("Migração do site atual", "Importação de notícias, documentos e estrutura, preservando SEO"),
 ("Treinamento e suporte", "Capacitação das equipes e suporte continuado"),
]:
    row = tb.add_row().cells
    set_cell_text(row[0], a, bold=True, size=9.5, color=AZUL_ESC); set_cell_text(row[1], b, size=9.5)

# ---------------- 12. ROADMAP ----------------
h1("Evolução contínua (roadmap)")
para("O produto evolui em fases com critérios de saída objetivos. As fases de Fundação, "
     "Transparência, Serviços+CMS, Diário Oficial e Inteligência Artificial estão concluídas e "
     "em produção. As próximas fases ampliam o App do Cidadão e a escala multi-região.")
img("roadmap", width=Inches(6.4))

# ---------------- 13. POR QUE LIDERA ----------------
h1("Por que a Lidera")
bullets([
 "Produto único que substitui vários contratos e fornecedores.",
 "Conformidade legal e acessibilidade no núcleo — não como adicional.",
 "Tecnologia moderna, segura e auditável, com IA integrada.",
 "Migração do site atual sem perder histórico nem posicionamento no Google.",
 "Implantação rápida, em nuvem ou no servidor da prefeitura.",
])

# ---------------- 14. CTA ----------------
tb = doc.add_table(rows=1, cols=1); no_table_borders(tb); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
c = tb.rows[0].cells[0]; shade(c, AZUL_ESC); c.text = ""
p1 = c.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.space_before = Pt(10)
r = p1.add_run("Vamos modernizar a sua prefeitura?"); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = HEX("#FFFFFF")
p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Agende uma demonstração gratuita e veja o portal funcionando com os dados do seu município.")
r2.font.size = Pt(10.5); r2.font.color.rgb = HEX("#E8EEFB")
p3 = c.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(10)
r3 = p3.add_run("Lidera Tecnologia e Gestão   ·   comercial@lideratecnologia.com.br   ·   www.lideratecnologia.com.br")
r3.font.size = Pt(10); r3.bold = True; r3.font.color.rgb = HEX(AMARELO)
para("Investimento sob consulta, conforme o porte do município.", size=9, italic=True,
     color="#667788").alignment = WD_ALIGN_PARAGRAPH.CENTER

out = os.path.join(BASE, "Portfolio-Portal-Prefeitura.docx")
doc.save(out)
print("OK ->", out)
